from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement, ns

document = Document()

header_lists = [0,1,2,3,4,5,6,7,8,9]
document.add_heading('Document Title', 0)

for x in header_lists:
    document.add_heading('Document Title', x)
    

print('word_test')
document.save('demo1.docx')